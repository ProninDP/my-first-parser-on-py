import requests
from bs4 import BeautifulSoup
from docx import Document
import re

def get_html(url): # забираем страницу
    r = requests.get(url) # проверяем доступ
    r.encoding = 'cp1251' # применяем нужную кодировку ибо вместо русского случаются крякозябры
    return r.text

def get_total_pages(html):  # список сылок
    soup = BeautifulSoup(html, 'lxml') # парсим страницу
    pages = soup.find_all('a') # ищем все с тегом "a"
    tmp = [link.get('href') for link in pages] # переводим во временный список ссылки
    output = [s for s in tmp if s.find('dialog_') != -1] # переводим в список ссылок которые начинаются на "dialog_"
    return output

def get_page_data(html):    # данные страниц
    try: # фильтруем ошибки доступа к атрибуту
        soup = BeautifulSoup(html, 'lxml') # парсим страницу
        pages = soup.find('table', class_='table').find_all('td')  # ищем все с тегом "td" в классе table
        data = [d.text for d in pages]    # создаем список данных, оставляем нужный нам текст
    except AttributeError: return False
    return data

def eng_rus(data): # разделяем диалоги на русские и английские
    e = re.compile("[a-zA-Z]+")
    r = re.compile("[а-яА-Я]+")
    eng = [w for w in filter(e.match, data)]
    rus = [w for w in filter(r.match, data)]
    return eng, rus # передаем кортеж

def new_document(dialog, en, ru): # генерируем новый word документ, печатаем данные
    document = Document()  # создаем новый word документ
    document.add_heading(dialog, 0)  # печатем загаловок с названием диалога
    table = document.add_table(rows=1, cols=2) # создаем таблицу
    hdr_cells = table.rows[0].cells # создаем заглавные ячейки таблицы
    hdr_cells[0].text = ru[0] # из диалогов на русском "По-английски" "название столбца"
    hdr_cells[1].text = ru[1] # из диалогов на русском "Перевод на русский" "название столбца"
    for a, b in zip(range(len(en)), range(2, len(ru))): # печатаем дилоги
        row_cells = table.add_row().cells # добавляем ячеки таблицы
        row_cells[0].text = en[a] # из диалогов на русском
        row_cells[1].text = ru[b] # из диалогов на английском
    document.save(dialog + '.docx')

def main():
    url = 'https://www.en365.ru/dialogi.htm' # страница с сылками на диалоги
    base_url = 'https://www.en365.ru/' # корневая ссылка

    out_url = get_total_pages(get_html(url)) # собираем список сылок(дилогов)

    try:
        for dialog in out_url: # перебираем список сылок(диалогов)
            url_gen = base_url + dialog # генерируем сылки с диалогами
            html = get_html(url_gen) # получаем данные по сылкам
            en, ru = eng_rus(get_page_data(html)) # собираем список данных со страниц и обрабатываем данные
            new_document(dialog, en, ru) # создаем новый word документ с полученными данными
    except TypeError: return False

if __name__ == '__main__':
    main()